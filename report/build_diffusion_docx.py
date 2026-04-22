"""Build the diffusion draft DOCX and embed the generated figure PNG files.

This script is intended to run in the ``mda`` environment, where ``python-docx``
is available. The HTML source is first converted with ``textutil`` so the
manuscript text and table layout remain consistent with the existing workflow.
The resulting DOCX is then updated with ``python-docx`` to insert the manuscript
figures directly into the package before their caption paragraphs.

Typical usage
-------------
1. ``textutil -convert docx report/diffusion_DM_source.html -output report/diffusion_DM.docx``
2. ``/Users/dm/miniforge3/envs/mda/bin/python3 report/build_diffusion_docx.py --skip-convert``

If ``textutil`` runs reliably in the current session, the full workflow can also
be executed directly via this script without ``--skip-convert``.
"""

from __future__ import annotations

import argparse
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class FigureInsertion:
    """Describe one figure to embed in the manuscript draft.

    Parameters
    ----------
    caption_prefix : str
        Leading text used to locate the target caption paragraph.
    image_path : pathlib.Path
        PNG file to insert before the caption.
    width_inches : float
        Display width of the inserted figure in inches.
    """

    caption_prefix: str
    image_path: Path
    width_inches: float


def parse_args() -> argparse.Namespace:
    """Parse command-line arguments for the document builder.

    Returns
    -------
    argparse.Namespace
        Parsed arguments describing the input source and output files.
    """

    report_dir = Path(__file__).resolve().parent
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source-html",
        type=Path,
        default=report_dir / "diffusion_DM_source.html",
        help="HTML source used for the draft text and captions.",
    )
    parser.add_argument(
        "--output-docx",
        type=Path,
        default=report_dir / "diffusion_DM.docx",
        help="DOCX file to create.",
    )
    parser.add_argument(
        "--output-rtf",
        type=Path,
        default=report_dir / "diffusion_DM_source.rtf",
        help="RTF file to create from the same HTML source.",
    )
    parser.add_argument(
        "--skip-convert",
        action="store_true",
        help="Skip the textutil conversion step and only embed figures into an existing DOCX.",
    )
    return parser.parse_args()


def run_textutil(source_html: Path, output_path: Path, *, convert_format: str) -> None:
    """Convert the HTML source into one output document with ``textutil``.

    Parameters
    ----------
    source_html : pathlib.Path
        HTML source file to convert.
    output_path : pathlib.Path
        Output path to create.
    convert_format : str
        ``textutil`` output format, e.g. ``"docx"`` or ``"rtf"``.
    """

    subprocess.run(
        [
            "textutil",
            "-convert",
            convert_format,
            str(source_html),
            "-output",
            str(output_path),
        ],
        check=True,
    )


def find_caption_paragraph(document: Document, caption_prefix: str) -> Paragraph:
    """Return the paragraph whose text starts with the requested caption prefix.

    Parameters
    ----------
    document : docx.document.Document
        Document to scan.
    caption_prefix : str
        Prefix used to identify the caption paragraph.

    Returns
    -------
    docx.text.paragraph.Paragraph
        Matching paragraph from the document.

    Raises
    ------
    ValueError
        If no paragraph starts with ``caption_prefix``.
    """

    for paragraph in document.paragraphs:
        if paragraph.text.startswith(caption_prefix):
            return paragraph
    raise ValueError(f"Could not find caption paragraph starting with: {caption_prefix!r}")


def build_figure_insertions(report_dir: Path) -> tuple[FigureInsertion, ...]:
    """Return the figures that should be embedded into the diffusion draft.

    Parameters
    ----------
    report_dir : pathlib.Path
        Directory containing the manuscript source and generated figures.

    Returns
    -------
    tuple[FigureInsertion, ...]
        Ordered figure definitions for the main-text and SI placeholders.
    """

    return (
        FigureInsertion(
            caption_prefix="Figure X.",
            image_path=report_dir / "diffusion_figures" / "diffusion_pooled_summary.png",
            width_inches=6.30,
        ),
        FigureInsertion(
            caption_prefix="Figure Y.",
            image_path=report_dir / "diffusion_figures" / "diffusion_ratio_summary.png",
            width_inches=6.10,
        ),
        FigureInsertion(
            caption_prefix="Figure S1.",
            image_path=report_dir / "diffusion_figures" / "diffusion_replica_summary.png",
            width_inches=6.30,
        ),
    )


def validate_figure_insertions(insertions: tuple[FigureInsertion, ...]) -> None:
    """Ensure that all requested figure files exist before editing the DOCX.

    Parameters
    ----------
    insertions : tuple[FigureInsertion, ...]
        Figure definitions that will be embedded into the document.

    Raises
    ------
    FileNotFoundError
        If any referenced figure file does not exist.
    """

    missing_paths = [str(insertion.image_path) for insertion in insertions if not insertion.image_path.is_file()]
    if missing_paths:
        missing_text = "\n".join(missing_paths)
        raise FileNotFoundError(f"Missing figure file(s):\n{missing_text}")


def remove_existing_figure_before_caption(caption_paragraph: Paragraph) -> None:
    """Remove an already-inserted figure paragraph immediately before a caption.

    This keeps the builder idempotent when ``--skip-convert`` is used on a DOCX
    that already contains embedded figures from an earlier run.

    Parameters
    ----------
    caption_paragraph : docx.text.paragraph.Paragraph
        Caption paragraph whose preceding sibling may contain a figure.
    """

    previous_element = caption_paragraph._p.getprevious()
    if previous_element is None:
        return
    if previous_element.xpath(".//w:drawing"):
        parent = previous_element.getparent()
        if parent is not None:
            parent.remove(previous_element)


def insert_figure_before_caption(document: Document, insertion: FigureInsertion) -> None:
    """Insert one figure immediately before its matching caption paragraph.

    Parameters
    ----------
    document : docx.document.Document
        Document being updated.
    insertion : FigureInsertion
        Figure metadata describing the caption target and image path.
    """

    caption_paragraph = find_caption_paragraph(document, insertion.caption_prefix)
    remove_existing_figure_before_caption(caption_paragraph)
    figure_paragraph = caption_paragraph.insert_paragraph_before()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.add_run().add_picture(
        str(insertion.image_path),
        width=Inches(insertion.width_inches),
    )


def build_document(
    source_html: Path,
    output_docx: Path,
    output_rtf: Path,
    *,
    skip_convert: bool,
) -> None:
    """Rebuild the draft documents and embed the current figures into the DOCX.

    Parameters
    ----------
    source_html : pathlib.Path
        HTML source used for text and captions.
    output_docx : pathlib.Path
        DOCX file to create and update.
    output_rtf : pathlib.Path
        RTF file to create from the same source.
    skip_convert : bool
        If ``True``, reuse the existing DOCX and RTF files instead of rerunning
        ``textutil``.
    """

    if not skip_convert:
        run_textutil(source_html, output_rtf, convert_format="rtf")
        run_textutil(source_html, output_docx, convert_format="docx")

    report_dir = source_html.parent
    insertions = build_figure_insertions(report_dir)
    validate_figure_insertions(insertions)

    document = Document(str(output_docx))
    for insertion in insertions:
        insert_figure_before_caption(document, insertion)
    document.save(str(output_docx))


def main() -> None:
    """Run the draft rebuild workflow."""

    args = parse_args()
    build_document(
        args.source_html,
        args.output_docx,
        args.output_rtf,
        skip_convert=args.skip_convert,
    )


if __name__ == "__main__":
    main()
